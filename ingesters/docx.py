from pathlib import Path
from ingesters import Document


def extract_docx(path: str | Path) -> Document:
    """Extract text content from a DOCX file."""
    from docx import Document as DocxDocument

    doc_path = Path(path)
    doc = DocxDocument(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return Document(raw_text=text, content_type="document")
